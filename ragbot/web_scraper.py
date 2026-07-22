"""
ragbot/web_scraper.py

Fetches and extracts text from web sources — both PDF and HTML.
Includes visual debugging (screenshots) and enhanced trace reporting.
"""
from __future__ import annotations

import csv as csv_mod
import hashlib
import io
import logging
import math
import random
import time
import tempfile
import os
import urllib.robotparser
from urllib.parse import urljoin, urlparse, urlunparse
from django.conf import settings

import requests

logger = logging.getLogger(__name__)
from filelock import FileLock
import os

_UC_DRIVER_LOCK_PATH = "/tmp/uc_driver_patch.lock"

REQUEST_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-AU,en;q=0.5",
}
REQUEST_TIMEOUT = 30

SKIP_EXTENSIONS = {".zip", ".png", ".jpg", ".jpeg", ".gif", ".svg", ".mp4", ".mp3"}

# Extensions we download and extract in-memory (never saved to disk).
# PDF is handled separately via Content-Type detection.
DOWNLOAD_EXTENSIONS = {".docx", ".doc", ".csv"}

# Extensions treated as "attachments" when discovered via links on an HTML
# page during sitemap-mode crawling (see _collect_attachment_links). Includes
# PDF in addition to DOWNLOAD_EXTENSIONS, since PDFs on the main crawl path
# are normally identified via Content-Type sniffing rather than extension,
# but attachment discovery happens before any request is made.
ATTACHMENT_EXTENSIONS = {".pdf", ".docx", ".doc", ".csv"}

# ---------------------------------------------------------------------------
# Human-behaviour constants
# ---------------------------------------------------------------------------
SESSION_ROTATE_EVERY = 40
DWELL_SHORT          = (3.0, 10.0)
DWELL_LONG           = (25.0, 75.0)
DWELL_LONG_PROB      = 0.15
BREAK_EVERY          = random.randint(18, 28)
BREAK_SLEEP          = (45.0, 120.0)
SCROLL_STEPS         = (2, 5)
SCROLL_AMOUNT        = (200, 600)

# ---------------------------------------------------------------------------
# Cloudflare back-off constants
# ---------------------------------------------------------------------------
# After this many CF challenges in one session, pause for a long time.
CF_BACKOFF_TRIGGER  = 3
CF_BACKOFF_SLEEP    = (300, 600)   # 5–10 minutes

CONTENT_SELECTORS = [
    "article", "main", ".legislation-content",
    "#content", ".content", "body",
]

SIDEBAR_SELECTORS = [
    "nav", "aside",
    '[class*="sidebar"]', '[class*="side-bar"]',
    '[class*="subnav"]',  '[class*="sub-nav"]',
    '[class*="related"]', '[class*="section-nav"]',
    '[id*="sidebar"]',    '[id*="subnav"]',
]

# --- VISUAL DEBUGGING ---
DEBUG_DIR = os.path.join(os.getcwd(), "scraper_debug")
os.makedirs(DEBUG_DIR, exist_ok=True)


# ---------------------------------------------------------------------------
# ★  DOMAIN CRAWL RULES  ★
# ---------------------------------------------------------------------------
# Single place for per-domain rules. Keys = domain netloc.
#
#   allowed_prefixes  – only crawl URLs starting with one of these.
#                       Empty = no restriction (all same-domain URLs allowed).
#   blocked_nav_ids   – HTML element IDs to strip before link discovery.
#   disallowed_paths  – path prefixes from robots.txt to never crawl.
#                       Supplemented at runtime by live robots.txt parsing.
# ---------------------------------------------------------------------------
DOMAIN_CRAWL_RULES: dict[str, dict] = {

    # ── acecqa.gov.au ─────────────────────────────────────────────────────────
    # Crawl-delay: 3 (from robots.txt)
    # ai-input permitted (Content-Signal: search=yes, ai-train=no)
    # Sitemap 403 via requests — fetched via Selenium at runtime
    "www.acecqa.gov.au": {
        "allowed_prefixes": [],
        "blocked_nav_ids":  [],
        "disallowed_paths": [
            "/core/", "/profiles/", "/admin/", "/comment/reply/",
            "/filter/tips", "/node/add/", "/search/", "/user/",
            "/book/export/html/", "/index.php/",
            "/README.txt", "/web.config",
        ],
        "sitemap_url": "https://www.acecqa.gov.au/sitemap.xml",
        "sitemap_needs_selenium": True,
    },

    # ── www.nsw.gov.au ────────────────────────────────────────────────────────
    # Sitemap index with 8 child sitemaps — accessible via requests
    "www.nsw.gov.au": {
        "allowed_prefixes": [],
        "blocked_nav_ids":  [],
        "disallowed_paths": [
            "/core/", "/profiles/", "/admin/", "/comment/reply/",
            "/filter/tips", "/node/add/", "/search/",
            "/user/register", "/user/password", "/user/login", "/user/logout",
            "/media/oembed", "/index.php/",
            "/archive/", "/covid-business/generate-qr-code/",
            "/image_popup/", "/node/", "/preview-link/",
            "/sites/default/files/noindex/", "/taxonomy/",
            "/tfa/", "/life-events/tfa/",
            "/form/nsw-transport-minister-contact",
            "/form/premier-contact-form",
            "/README.txt", "/web.config",
        ],
        "sitemap_url": "https://www.nsw.gov.au/sitemap.xml",
        "sitemap_needs_selenium": True,
    },

    # ── www.education.gov.au ─────────────────────────────────────────────────
    # Only crawl /early-childhood subtree (client requirement)
    # Sitemap times out via requests — fetched via Selenium
    "www.education.gov.au": {
        "allowed_prefixes": [],
        "blocked_nav_ids":  [],
        "disallowed_paths": [
            "/admin/", "/user/", "/search/", "/node/",
            "/comment/reply/", "/filter/tips",
        ],
        "sitemap_url": "https://www.education.gov.au/sitemap.xml",
        "sitemap_needs_selenium": True,
    },

    # ── education.nsw.gov.au ─────────────────────────────────────────────────
    # Only crawl /early-childhood-education subtree (client requirement)
    # Sitemap index with 43 child sitemaps — accessible via requests
    "education.nsw.gov.au": {
        "allowed_prefixes": [],
        "blocked_nav_ids":  [],
        "disallowed_paths": [
            "/admin/", "/user/", "/search/", "/node/",
            "/comment/reply/", "/filter/tips",
        ],
        "sitemap_url": "https://education.nsw.gov.au/sitemap.xml",
        "sitemap_needs_selenium": True,
    },

    # ── legislation.nsw.gov.au ───────────────────────────────────────────────
    # NSW legislation portal — rich compliance content
    # Non-standard sitemap path: /sitemaps/sitemapindex.xml
    "legislation.nsw.gov.au": {
        "allowed_prefixes": [],
        "blocked_nav_ids":  [],
        "disallowed_paths": [
            "/search/", "/DownloadFile/", "/printview/",
            "/admin/", "/user/",
        ],
        "sitemap_url": "https://legislation.nsw.gov.au/sitemaps/sitemapindex.xml",
        "sitemap_needs_selenium": True,
    },
}


def _get_domain_rules(url: str) -> dict:
    netloc = urlparse(url).netloc
    return DOMAIN_CRAWL_RULES.get(
        netloc,
        {"allowed_prefixes": [], "blocked_nav_ids": [], "disallowed_paths": []},
    )


# ---------------------------------------------------------------------------
# robots.txt live fetching
# ---------------------------------------------------------------------------

def _fetch_robots_rules(root_url: str) -> dict:
    """
    Fetch and parse robots.txt for root_url's domain at crawl start.
    Returns:
        crawl_delay  – seconds to wait between requests (float)
        parser       – RobotFileParser instance for can_fetch() checks per URL
    Falls back to crawl_delay=3.0 and parser=None on any error.

    Deliberately does NOT use RobotFileParser.read() directly: that method
    fetches robots.txt via bare urllib.request.urlopen() with no User-Agent
    header, and a WAF-protected site returning 403 to that bare request
    causes CPython's robotparser to set disallow_all=True — which makes
    can_fetch() return False for every single URL on the domain for the
    rest of the run, silently zeroing out an otherwise-valid sitemap. We
    fetch the text ourselves with browser-like REQUEST_HEADERS (same as
    every other request in this file) and feed it to rp.parse() instead.
    """
    parsed     = urlparse(root_url)
    robots_url = f"{parsed.scheme}://{parsed.netloc}/robots.txt"
    rp         = urllib.robotparser.RobotFileParser()
    rp.set_url(robots_url)
    try:
        sess = requests.Session()
        sess.headers.update(REQUEST_HEADERS)
        resp = sess.get(robots_url, timeout=10)

        if resp.status_code == 200:
            rp.parse(resp.text.splitlines())
        elif 400 <= resp.status_code < 500:
            # Covers 401/403 too: a WAF/permission response on robots.txt
            # itself is not a site-wide Disallow directive. Treat as "no
            # live rules found" rather than "disallow everything" — the
            # hardcoded disallowed_paths in DOMAIN_CRAWL_RULES still apply.
            logger.warning(
                f"robots.txt at {robots_url} returned {resp.status_code} — "
                f"treating as no live rules (not disallow-all)."
            )
            rp = None
        else:
            logger.warning(f"robots.txt at {robots_url} returned {resp.status_code}")
            rp = None

        delay = 3.0
        if rp is not None:
            delay = (
                rp.crawl_delay("Mozilla")
                or rp.crawl_delay("*")
                or 3.0
            )
        print(f"🤖 robots.txt fetched from {robots_url} — Crawl-delay: {delay}s")
        return {"crawl_delay": float(delay), "parser": rp}
    except Exception as exc:
        logger.warning(f"Could not fetch robots.txt from {robots_url}: {exc}")
        return {"crawl_delay": 3.0, "parser": None}


# ---------------------------------------------------------------------------
# URL filtering helpers
# ---------------------------------------------------------------------------

def _normalize_sitemap_url(url: str) -> str:
    """
    Domain-specific URL normalization applied to sitemap entries before
    filtering/dedup.

    nsw.gov.au (and any subdomain of it) specific: sitemap
    entries sometimes point at /view/html/... with a fragment identifying
    a specific schedule/section (e.g. #sch.6-sec.34). That fragment targets
    an anchor inside the partial "html" view. We rewrite these to the
    corresponding /view/whole/html/... URL — the complete act/instrument in
    a single page — and drop the fragment entirely, so the crawl fetches
    one canonical full-document URL per act instead of scraping/duplicating
    the same document once per provision-level anchor.
    Only applies to nsw.gov.au and its subdomains; every other
    domain's URLs pass through unchanged (fragment stripping for those still
    happens via the "#" check in _is_allowed_href / the split("#")[0] callers).
    """
    netloc = urlparse(url).netloc
    if netloc != "nsw.gov.au" and not netloc.endswith(".nsw.gov.au"):
        return url

    parsed = urlparse(url)
    path   = parsed.path
    if path.startswith("/view/html/"):
        new_path = "/view/whole/html/" + path[len("/view/html/"):]
        parsed   = parsed._replace(path=new_path, fragment="")
        url      = urlunparse(parsed)
    else:
        # No /view/html/ prefix to rewrite — still drop any fragment.
        parsed = parsed._replace(fragment="")
        url    = urlunparse(parsed)
    return url


# ---------------------------------------------------------------------------
# ★  SITEMAP DISCOVERY & PARSING  ★
# ---------------------------------------------------------------------------

SITEMAP_PROBE_PATHS = [
    "/sitemap.xml",
    "/sitemap_index.xml",
    "/sitemap-index.xml",
    "/sitemaps/sitemapindex.xml",
    "/sitemap/sitemap.xml",
]


def _fetch_sitemap_xml(url: str, use_selenium: bool = False, driver=None) -> str | None:
    """
    Fetch raw XML from a sitemap URL.
    Tries requests first; falls back to Selenium if use_selenium=True.
    Reuses existing driver if provided, avoiding cold session WAF blocks on child sitemaps.
    """
    # Try requests first (fast)
    try:
        sess = requests.Session()
        sess.headers.update(REQUEST_HEADERS)
        resp = sess.get(url, timeout=15, allow_redirects=True)
        if resp.status_code == 200 and (
            "xml" in resp.headers.get("Content-Type", "")
            or resp.text.strip().startswith("<?xml")
            or "<urlset" in resp.text
            or "<sitemapindex" in resp.text
        ):
            return resp.text
    except Exception:
        pass

    if not use_selenium:
        return None

    # Selenium fallback (for 403/CF-protected sitemaps)
    local_driver = False
    if driver is None:
        driver = _make_selenium_driver()
        local_driver = True

    try:
        print(f"🌍 [Selenium] Fetching sitemap: {url}")
        driver.get(url)
        _wait_for_cloudflare_if_needed(driver)
        _wait_for_page_ready(driver)
        _capture_screenshot(driver, "sitemap_fetch")

        # Strategy 1: in-browser fetch() — gets raw XML before Chrome wraps it in HTML
        xml_text = None
        try:
            xml_text = driver.execute_async_script(
                "const cb = arguments[arguments.length-1];"
                "fetch(arguments[0])"
                ".then(r => r.text())"
                ".then(t => cb(t))"
                ".catch(() => cb(''))",
                url,
            )
        except Exception:
            pass

        # Strategy 2: innerText/textContent — Chrome decodes HTML entities
        if not xml_text or not ("<urlset" in xml_text or "<sitemapindex" in xml_text):
            try:
                xml_text = driver.execute_script(
                    "return document.body.innerText || document.body.textContent || ''"
                )
            except Exception:
                xml_text = None

        # Strategy 3: raw page_source — last resort, may have HTML-escaped XML
        if not xml_text or not ("<urlset" in xml_text or "<sitemapindex" in xml_text):
            xml_text = driver.page_source or ""
            # Unescape HTML entities Chrome may have encoded
            xml_text = (xml_text
                .replace("&lt;", "<")
                .replace("&gt;", ">")
                .replace("&amp;", "&")
                .replace("&quot;", '"')
            )

        if "<urlset" in xml_text or "<sitemapindex" in xml_text:
            # Extract just the XML portion
            import re as _re
            match = _re.search(
                r"(<(?:urlset|sitemapindex)[\s\S]*?</(?:urlset|sitemapindex)>)",
                xml_text,
                _re.DOTALL,
            )
            return match.group(1) if match else xml_text

        logger.warning(f"Selenium sitemap fetch: no XML found at {url}")
        return None
    except Exception as exc:
        logger.warning(f"Selenium sitemap fetch failed {url}: {exc}")
        return None
    finally:
        if local_driver and driver:
            driver.quit()


def _parse_sitemap_xml(xml: str, root_url: str, robots_parser=None, driver=None) -> list[str]:
    """
    Parse a sitemap or sitemap index XML string.
    Recursively fetches child sitemaps if it's an index.
    Filters URLs through domain rules and robots.txt.
    Returns a flat deduplicated list of page URLs.
    """
    import xml.etree.ElementTree as ET
    import re

    def _clean_xml(xml_str: str) -> str:
        # Step 1: unescape HTML entities Selenium may have encoded
        xml_str = (xml_str
            .replace("&lt;", "<").replace("&gt;", ">")
            .replace("&amp;", "&").replace("&quot;", '"')
        )
        # Step 2: extract just the XML block — strips Chrome HTML wrapper
        for tag in ("sitemapindex", "urlset"):
            match = re.search(
                rf"(<{tag}[\s\S]*?</{tag}>)",
                xml_str,
                re.DOTALL | re.IGNORECASE,
            )
            if match:
                xml_str = match.group(1)
                break
        # Step 3: remove ALL xmlns declarations
        xml_str = re.sub(r'\s+xmlns(?::[a-zA-Z0-9_]+)?="[^"]*"', "", xml_str)
        xml_str = re.sub(r"\s+xmlns(?::[a-zA-Z0-9_]+)?='[^']*'", "", xml_str)
        # Step 4: remove prefixed ATTRIBUTES
        xml_str = re.sub(r'\s+[a-zA-Z0-9_]+:[a-zA-Z0-9_]+=(?:"[^"]*"|\'[^\']*\')', "", xml_str)
        # Step 5: remove namespace prefixes from TAG names
        xml_str = re.sub(r"<(/?)([a-zA-Z0-9_]+):([a-zA-Z0-9_]+)", r"<\1\3", xml_str)
        # Step 6: strip XML declaration
        xml_str = re.sub(r"<\?xml[^?]*\?>", "", xml_str)
        return xml_str.strip()

    def _extract_locs_regex(xml_str: str) -> list[str]:
        """Last-resort extraction: pull <loc> values with regex, no XML parse."""
        xml_unescaped = (xml_str
            .replace("&lt;", "<").replace("&gt;", ">")
            .replace("&amp;", "&").replace("&quot;", '"')
        )
        # Added re.IGNORECASE because Chrome page_source sometimes capitalizes tags
        locs = re.findall(r"<loc>\s*(.*?)\s*</loc>", xml_unescaped, re.DOTALL | re.IGNORECASE)
        if not locs:
            locs = re.findall(r"https?://[^\s<>\"']+", xml_unescaped)
        return [m.strip() for m in locs if m.strip()]

    # Try ET parse with namespace stripping
    xml_clean = _clean_xml(xml)
    root = None
    try:
        root = ET.fromstring(xml_clean)
    except ET.ParseError as e:
        logger.warning(f"Sitemap XML parse error after namespace strip: {e} — trying regex fallback")

    urls = []
    rules          = _get_domain_rules(root_url)
    needs_selenium = rules.get("sitemap_needs_selenium", False)

    # ── Sitemap index ─────────────────────────────────────────────────────
    is_index = (
        (root is not None and (root.tag == "sitemapindex" or root.tag.endswith("}sitemapindex")))
        or "<sitemapindex" in xml
    )
    if is_index:
        if root is not None:
            child_urls = [loc.text.strip() for loc in root.findall(".//loc") if loc.text]
            # Force fallback if namespace issues hid the <loc> tags from ET
            if not child_urls:
                child_urls = _extract_locs_regex(xml)
        else:
            child_urls = _extract_locs_regex(xml)
            
        print(f"  📑 Sitemap index: {len(child_urls)} child sitemaps extracted")
        for child_url in child_urls:
            child_url = child_url.strip()
            if not child_url:
                continue
            print(f"    → Fetching child sitemap: {child_url}")
            child_xml = _fetch_sitemap_xml(child_url, use_selenium=needs_selenium, driver=driver)
            
            if child_xml:
                child_urls_parsed = _parse_sitemap_xml(child_xml, root_url, robots_parser, driver=driver)
                urls.extend(child_urls_parsed)
        return urls

    # ── Regular sitemap ───────────────────────────────────────────────────
    if root is not None:
        raw_urls = [loc.text.strip() for loc in root.findall(".//loc") if loc.text]
        # Force fallback if namespace issues hid the <loc> tags from ET
        if not raw_urls:
            raw_urls = _extract_locs_regex(xml)
            if raw_urls:
                print("  ℹ️  Used regex fallback because ElementTree found 0 <loc> tags.")
    else:
        raw_urls = _extract_locs_regex(xml)
        logger.info(f"  ℹ️  Used regex fallback — extracted {len(raw_urls)} raw URLs")

    print(f"  🔍 Extracted {len(raw_urls)} raw URLs from XML before filtering")

    filtered_count = 0
    for url in raw_urls:
        url = _normalize_sitemap_url(url)
        url = url.rstrip("/")
        if not url:
            continue
            
        if _is_allowed_href(url, root_url, robots_parser):
            urls.append(url)
        else:
            filtered_count += 1
            # Temporarily print the first 3 filtered URLs so we can see WHY they are blocked
            if filtered_count <= 3:
                print(f"  🚫 Filtered out: {url}")
            logger.debug(f"  🚫 Sitemap URL filtered: {url}")

    if filtered_count:
        print(f"  ℹ️  {filtered_count} sitemap URLs filtered by domain rules/robots.txt")
        logger.info(f"  ℹ️  {filtered_count} sitemap URLs filtered by domain rules/robots.txt")
        
    print(f"  📄 {len(urls)} URLs fully extracted and added to queue from this sitemap")

    return urls
def _discover_sitemap(root_url: str, robots_parser=None) -> list[str] | None:
    """
    Discover and parse the sitemap for root_url.

    Strategy (in order):
      1. Check DOMAIN_CRAWL_RULES for a known sitemap_url.
      2. Check robots.txt for a Sitemap: directive.
      3. Probe standard sitemap paths (/sitemap.xml etc).
      4. Fetch homepage and look for <link rel="sitemap"> in <head>.
      5. Return None if no sitemap found (caller falls back to BFS crawl).

    Returns a flat deduplicated list of filtered page URLs, or None.
    """
    rules          = _get_domain_rules(root_url)
    needs_selenium = rules.get("sitemap_needs_selenium", False)

    driver = None
    if needs_selenium:
        driver = _make_selenium_driver()

    try:
        # ── 1. Known sitemap from DOMAIN_CRAWL_RULES ─────────────────────────
        known = rules.get("sitemap_url")
        if known:
            print(f"🗺️  Known sitemap: {known}")
            xml = _fetch_sitemap_xml(known, use_selenium=needs_selenium, driver=driver)
            if xml:
                urls = _parse_sitemap_xml(xml, root_url, robots_parser, driver=driver)
                if urls:
                    print(f"  ✅ {len(urls)} URLs from known sitemap")
                    return list(dict.fromkeys(urls))
                else:
                    # XML fetched OK but all URLs filtered — log clearly and stop
                    # Don't fall through to probe: the sitemap is the right one,
                    # the issue is filtering rules (e.g. allowed_prefixes too strict).
                    logger.warning(
                        f"  ⚠️  Known sitemap {known} fetched OK but 0 URLs passed "
                        f"domain filters. Check allowed_prefixes in DOMAIN_CRAWL_RULES "
                        f"or disallowed_paths. Falling back to BFS crawl."
                    )
                    return None
            else:
                logger.warning(f"  ⚠️  Known sitemap {known} could not be fetched — trying other methods.")

        # ── 2. robots.txt Sitemap: directive ─────────────────────────────────
        if robots_parser:
            # urllib.robotparser doesn't expose Sitemap directives directly,
            # so re-fetch and parse manually
            parsed     = urlparse(root_url)
            robots_url = f"{parsed.scheme}://{parsed.netloc}/robots.txt"
            try:
                sess = requests.Session()
                sess.headers.update(REQUEST_HEADERS)
                resp = sess.get(robots_url, timeout=10)
                if resp.status_code == 200:
                    import re
                    sitemaps = re.findall(r"(?i)^Sitemap:\s*(.+)$", resp.text, re.MULTILINE)
                    for sm_url in sitemaps:
                        sm_url = sm_url.strip()
                        print(f"🗺️  Sitemap from robots.txt: {sm_url}")
                        xml = _fetch_sitemap_xml(sm_url, use_selenium=needs_selenium, driver=driver)
                        if xml:
                            urls = _parse_sitemap_xml(xml, root_url, robots_parser, driver=driver)
                            if urls:
                                print(f"  ✅ {len(urls)} URLs from robots.txt sitemap")
                                return list(dict.fromkeys(urls))
            except Exception as exc:
                logger.debug(f"robots.txt sitemap check failed: {exc}")

        # ── 3. Probe standard paths ───────────────────────────────────────────
        parsed = urlparse(root_url)
        base   = f"{parsed.scheme}://{parsed.netloc}"
        for path in SITEMAP_PROBE_PATHS:
            probe_url = base + path
            print(f"🔍 Probing: {probe_url}")
            xml = _fetch_sitemap_xml(probe_url, use_selenium=needs_selenium, driver=driver)
            if xml:
                urls = _parse_sitemap_xml(xml, root_url, robots_parser, driver=driver)
                if urls:
                    print(f"  ✅ {len(urls)} URLs from {path}")
                    return list(dict.fromkeys(urls))

        # ── 4. Homepage <link rel="sitemap"> ─────────────────────────────────
        try:
            from bs4 import BeautifulSoup
            sess = requests.Session()
            sess.headers.update(REQUEST_HEADERS)
            resp = sess.get(root_url, timeout=15)
            if resp.status_code == 200:
                soup = BeautifulSoup(resp.text, "html.parser")
                link = soup.find("link", rel="sitemap")
                if link and link.get("href"):
                    sm_url = urljoin(root_url, link["href"])
                    print(f"🗺️  Sitemap from homepage <link>: {sm_url}")
                    xml = _fetch_sitemap_xml(sm_url, use_selenium=needs_selenium, driver=driver)
                    if xml:
                        urls = _parse_sitemap_xml(xml, root_url, robots_parser, driver=driver)
                        if urls:
                            print(f"  ✅ {len(urls)} URLs from homepage sitemap link")
                            return list(dict.fromkeys(urls))
        except Exception as exc:
            logger.debug(f"Homepage sitemap link check failed: {exc}")

        print(f"⚠️  No sitemap found for {root_url} — will use BFS crawl")
        return None

    finally:
        if driver:
            driver.quit()


# ---------------------------------------------------------------------------
# Sitemap-based page fetcher (replaces BFS when sitemap is available)
# ---------------------------------------------------------------------------

def _crawl_from_sitemap(
    session: requests.Session,
    urls: list[str],
    root_url: str,
    use_selenium: bool = False,
) -> list[dict]:
    """
    Fetch and extract content from a pre-built list of URLs (from sitemap).
    Same extraction logic as _crawl_html but no link discovery needed —
    the sitemap already gives us the complete URL list.
    No max_pages limit — processes every URL in the sitemap.
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError("pip install beautifulsoup4 lxml")

    # Fetch robots.txt for crawl-delay
    robots        = _fetch_robots_rules(root_url)
    crawl_delay   = robots["crawl_delay"]
    robots_parser = robots["parser"]

    results:  list[dict] = []
    visited:  set[str]   = set()
    total                = len(urls)

    driver             = _make_selenium_driver() if use_selenium else None
    pages_fetched      = 0
    session_page_count = 0
    cf_challenge_count = 0

    print(f"🗺️  Sitemap crawl: {total} URLs to fetch from {root_url}")

    try:
        for url in urls:
            url = url.rstrip("/")
            if not url or url in visited or _should_skip(url):
                continue

            # ScrapedURL cache check
            try:
                from ragbot.models import ScrapedURL
                if ScrapedURL.was_scraped_today(url):
                    logger.info(f"⏭️  Already scraped today: {url}")
                    visited.add(url)
                    continue
            except Exception:
                pass

            visited.add(url)
            print(f"🔍 [{pages_fetched+1}/{total}] {url}")

            html_source = None
            final_url   = url
            page_hash   = None

            # Session rotation
            if use_selenium and driver and session_page_count >= SESSION_ROTATE_EVERY:
                driver = _restart_selenium_session(driver)
                session_page_count = 0
                cf_challenge_count = 0

            if use_selenium:
                try:
                    title, text, final_url, soup = _fetch_html_via_selenium(driver, url)

                    if not text and not title:
                        cf_challenge_count += 1
                        if cf_challenge_count >= CF_BACKOFF_TRIGGER:
                            backoff = random.uniform(*CF_BACKOFF_SLEEP)
                            print(f"🛑 CF {cf_challenge_count}x — backing off {backoff:.0f}s")
                            time.sleep(backoff)
                            cf_challenge_count = 0
                    else:
                        cf_challenge_count = 0

                    if text and text.strip():
                        page_hash = hashlib.sha256(text.encode()).hexdigest()
                        results.append({
                            "url":          final_url,
                            "title":        title[:500],
                            "text":         text,
                            "content_hash": page_hash,
                            "parent_url":   None,   # sitemap crawl — flat, no parent tracking
                        })
                        _discover_and_extract_attachments(
                            soup, final_url, root_url, robots_parser,
                            session, driver, visited, results,
                        )
                    pages_fetched      += 1
                    session_page_count += 1

                except _DocumentRedirect as doc_redir:
                    results.extend(doc_redir.pages)
                    if doc_redir.pages:
                        page_hash = doc_redir.pages[0]["content_hash"]
                    final_url          = doc_redir.final_url
                    pages_fetched      += 1
                    session_page_count += 1

                except Exception as exc:
                    logger.warning(f"Selenium failed for {url}: {exc}")
                    continue

            else:
                try:
                    resp = session.get(url, timeout=REQUEST_TIMEOUT)

                    if resp.status_code == 403:
                        if driver is None:
                            driver = _make_selenium_driver()
                            session_page_count = 0
                        try:
                            title, text, final_url, soup = _fetch_html_via_selenium(
                                driver, url, try_click=False
                            )
                            if text and text.strip():
                                page_hash = hashlib.sha256(text.encode()).hexdigest()
                                results.append({
                                    "url":          final_url,
                                    "title":        title[:500],
                                    "text":         text,
                                    "content_hash": page_hash,
                                    "parent_url":   None,
                                })
                                _discover_and_extract_attachments(
                                    soup, final_url, root_url, robots_parser,
                                    session, driver, visited, results,
                                )
                            pages_fetched      += 1
                            session_page_count += 1
                        except _DocumentRedirect as doc_redir:
                            results.extend(doc_redir.pages)
                            if doc_redir.pages:
                                page_hash = doc_redir.pages[0]["content_hash"]
                            final_url = doc_redir.final_url
                            pages_fetched += 1
                        except Exception as exc2:
                            logger.warning(f"Selenium fallback failed {url}: {exc2}")
                        continue

                    resp.raise_for_status()
                    ct = resp.headers.get("Content-Type", "")

                    if "application/pdf" in ct:
                        pdf_pages = _extract_pdf_pages(io.BytesIO(resp.content), url)
                        results.extend(pdf_pages)
                        if pdf_pages:
                            page_hash = pdf_pages[0]["content_hash"]
                        pages_fetched += 1
                        _human_delay_between_pages(pages_fetched, crawl_delay)
                        try:
                            ScrapedURL.mark_scraped(url, content_hash=page_hash or "")
                        except Exception:
                            pass
                        continue

                    if any(x in ct for x in [
                        "application/vnd.openxmlformats-officedocument.wordprocessingml",
                        "application/msword",
                    ]) or url.lower().endswith((".docx", ".doc")):
                        docx_pages = _extract_docx_bytes(io.BytesIO(resp.content), url)
                        results.extend(docx_pages)
                        if docx_pages:
                            page_hash = docx_pages[0]["content_hash"]
                        pages_fetched += 1
                        _human_delay_between_pages(pages_fetched, crawl_delay)
                        try:
                            ScrapedURL.mark_scraped(url, content_hash=page_hash or "")
                        except Exception:
                            pass
                        continue

                    if "text/csv" in ct or url.lower().endswith(".csv"):
                        csv_pages = _extract_csv_bytes(resp.content, url)
                        results.extend(csv_pages)
                        if csv_pages:
                            page_hash = csv_pages[0]["content_hash"]
                        pages_fetched += 1
                        _human_delay_between_pages(pages_fetched, crawl_delay)
                        try:
                            ScrapedURL.mark_scraped(url, content_hash=page_hash or "")
                        except Exception:
                            pass
                        continue

                    if "text/html" not in ct:
                        continue

                    soup  = BeautifulSoup(resp.text, "html.parser")
                    title = soup.title.string.strip() if soup.title and soup.title.string else url
                    text  = _extract_html_text(soup)
                    if text.strip():
                        page_hash = hashlib.sha256(text.encode()).hexdigest()
                        results.append({
                            "url":          url,
                            "title":        title[:500],
                            "text":         text,
                            "content_hash": page_hash,
                            "parent_url":   None,
                        })
                        _discover_and_extract_attachments(
                            soup, url, root_url, robots_parser,
                            session, driver, visited, results,
                        )
                    pages_fetched += 1

                except requests.RequestException as exc:
                    logger.warning(f"Request failed {url}: {exc}")
                    continue

            # Mark as scraped
            if page_hash:
                try:
                    ScrapedURL.mark_scraped(final_url, content_hash=page_hash)
                    if final_url != url:
                        ScrapedURL.mark_scraped(url, content_hash=page_hash)
                except Exception:
                    pass

            visited.add(final_url)
            _human_delay_between_pages(pages_fetched, crawl_delay)

    finally:
        if driver:
            print("🔌 Closing Selenium session.")
            driver.quit()

    print(f"✅ Sitemap crawl complete. {len(results)} pages from {root_url}")
    return results


# ---------------------------------------------------------------------------
# Attachment discovery (sitemap mode only)
# ---------------------------------------------------------------------------
# _crawl_from_sitemap only ever fetches the exact URLs handed to it by the
# sitemap — it does no link discovery of its own. That's a problem for pages
# like grant/guideline landing pages that link out to PDF/DOCX/CSV
# attachments that are NOT independently listed in the site's sitemap.xml
# (common on Drupal and similar CMSs, whose sitemap generators usually index
# content nodes, not raw file assets). The functions below scan each fetched
# HTML page for such attachment links and fetch/extract them too, tagging
# each with parent_url so the parent/child SourceDocument relationship is
# preserved.

def _collect_attachment_links(
    soup,
    current_url: str,
    root_url: str,
    robots_parser=None,
) -> list[str]:
    """
    Scan a page's <a href> tags for attachment links (PDF/DOCX/DOC/CSV).
    Reuses _is_allowed_href so the same domain/robots.txt/disallowed_paths
    rules apply as everywhere else in the crawler.
    """
    from bs4 import BeautifulSoup

    soup_copy = BeautifulSoup(str(soup), "html.parser")
    seen  = set()
    links = []
    for a in soup_copy.find_all("a", href=True):
        raw = a["href"].split("#")[0].strip()
        if not raw:
            continue
        href = urljoin(current_url, raw).rstrip("/")
        path = urlparse(href).path.lower()
        if not any(path.endswith(ext) for ext in ATTACHMENT_EXTENSIONS):
            continue
        if href in seen:
            continue
        seen.add(href)
        if _is_allowed_href(href, root_url, robots_parser):
            links.append(href)
    return links


def _fetch_attachment(session: requests.Session, url: str, driver=None) -> list[dict]:
    """
    Download and extract a single attachment discovered via
    _collect_attachment_links. Tries plain requests first — attachment files
    served from static file paths (e.g. Drupal's /sites/default/files/) are
    often not behind the same WAF as the parent page. Falls back to the
    existing Selenium document-redirect handling in _fetch_html_via_selenium
    if requests is blocked and a driver is available (this only covers
    PDF/DOCX/DOC — that function has no CSV-redirect detection).
    """
    path = urlparse(url).path.lower()
    try:
        resp = session.get(url, timeout=REQUEST_TIMEOUT)
        if resp.status_code == 200:
            if path.endswith(".pdf"):
                return _extract_pdf_pages(io.BytesIO(resp.content), url)
            if path.endswith((".docx", ".doc")):
                return _extract_docx_bytes(io.BytesIO(resp.content), url)
            if path.endswith(".csv"):
                return _extract_csv_bytes(resp.content, url)
            return []
        if resp.status_code != 403 or driver is None:
            logger.warning(f"Attachment fetch got {resp.status_code}: {url}")
            return []
    except Exception as exc:
        if driver is None:
            logger.warning(f"Attachment fetch failed {url}: {exc}")
            return []

    # 403 (or a request error) with a driver available — try Selenium's
    # existing PDF/DOCX redirect-detection path.
    try:
        _fetch_html_via_selenium(driver, url, try_click=False)
        # Reached only if the driver didn't recognize this as a document —
        # nothing more we can do with it.
        return []
    except _DocumentRedirect as doc_redir:
        return doc_redir.pages
    except Exception as exc:
        logger.warning(f"Attachment fetch failed via Selenium {url}: {exc}")
        return []


def _discover_and_extract_attachments(
    soup,
    parent_url: str,
    root_url: str,
    robots_parser,
    session: requests.Session,
    driver,
    visited: set,
    results: list,
) -> None:
    """
    Find and fetch any attachment links on a sitemap-crawled HTML page.
    Mutates `results` (extends) and `visited` (adds fetched attachment URLs)
    in place. Each extracted attachment page is tagged with parent_url.
    """
    if not soup:
        return
    for att_url in _collect_attachment_links(soup, parent_url, root_url, robots_parser):
        if att_url in visited:
            continue
        visited.add(att_url)

        try:
            from ragbot.models import ScrapedURL
            if ScrapedURL.was_scraped_today(att_url):
                logger.info(f"⏭️  Attachment already scraped today: {att_url}")
                continue
        except Exception:
            pass

        print(f"  📎 Attachment discovered on {parent_url}: {att_url}")
        pages = _fetch_attachment(session, att_url, driver=driver)
        for page in pages:
            page["parent_url"] = parent_url
        results.extend(pages)

        if pages:
            try:
                ScrapedURL.mark_scraped(att_url, content_hash=pages[0]["content_hash"])
            except Exception:
                pass

def _is_downloadable(url: str) -> bool:
    path = urlparse(url).path.lower()
    return any(path.endswith(ext) for ext in DOWNLOAD_EXTENSIONS)


def _same_domain(url: str, root: str) -> bool:
    return urlparse(url).netloc == urlparse(root).netloc


def _should_skip(url: str) -> bool:
    path = urlparse(url).path.lower()
    return any(path.endswith(ext) for ext in SKIP_EXTENSIONS)


def _normalize_crawl_url(url: str) -> str:
    """
    Strip URL fragments before filtering/dedup, with one domain-specific
    rewrite for legislation.nsw.gov.au (and any subdomain of it).

    That site serves provision-level deep links as #fragments on its
    "current view" page, e.g.:
        .../view/html/inforce/2001-12-01/act-1992-018#sch.6-sec.34
    Simply discarding the fragment there would leave a URL that's a
    real, fetchable page — but not the canonical whole-document URL.
    The site's actual single-page rendering of the full document lives
    at the equivalent .../view/whole/html/... path. So for this domain
    we rewrite /view/html/ -> /view/whole/html/ and drop the fragment,
    instead of just stripping it. This also means many distinct
    provision-fragment links collapse onto one canonical document URL,
    which is the correct de-dup behaviour (one document, many anchors).

    Every other domain just gets the fragment stripped, unchanged from
    prior behaviour.
    """
    if "#" not in url:
        return url

    base_url = url.split("#", 1)[0]
    if not base_url:
        # Pure in-page anchor (e.g. href="#sch.6-sec.34") with nothing
        # before the "#" — caller must supply the page URL to resolve
        # against (handled via urljoin before this function is called).
        return base_url

    parsed = urlparse(base_url)
    netloc = parsed.netloc

    if netloc == "legislation.nsw.gov.au" or netloc.endswith(".legislation.nsw.gov.au"):
        if "/view/html/" in parsed.path and "/view/whole/html/" not in parsed.path:
            new_path = parsed.path.replace("/view/html/", "/view/whole/html/", 1)
            base_url = f"{parsed.scheme}://{netloc}{new_path}"

    return base_url


def _is_allowed_href(href: str, root_url: str, robots_parser=None) -> bool:
    """
    Apply domain rules + live robots.txt to a candidate href.

    Order:
      0. Root URL itself — always allowed (prevents prefix rules from
         blocking the entry point of the crawl).
      1. Downloadable files (.docx/.csv) — allowed through immediately
         BUT must still be same-domain.
      2. Skip extensions (.zip/.png etc) — reject.
      3. Must be same domain as root_url.
      4. Fragment-only links — reject.
      5. robots.txt can_fetch() check (if parser available).
      6. Hardcoded disallowed_paths from DOMAIN_CRAWL_RULES.
      7. allowed_prefixes restriction (if set).
    """
    if not href:
        return False

    # Always allow the root URL itself through — prefix rules must not
    # block the crawl entry point (e.g. education.gov.au root vs
    # /early-childhood prefix restriction).
    if href.rstrip("/") == root_url.rstrip("/"):
        return True

    # Downloadable — same-domain check still applies
    if _is_downloadable(href):
        return _same_domain(href, root_url)

    if _should_skip(href):
        return False
    if not _same_domain(href, root_url):
        return False
    if "#" in href:
        return False

    # Live robots.txt check
    if robots_parser is not None:
        if not robots_parser.can_fetch("*", href):
            logger.debug(f"  🤖 robots.txt disallows: {href}")
            return False

    rules       = _get_domain_rules(href)
    parsed_path = urlparse(href).path

    # Hardcoded disallowed paths
    for blocked in rules.get("disallowed_paths", []):
        if parsed_path.startswith(blocked) or parsed_path == blocked.rstrip("/"):
            return False

    prefixes = rules.get("allowed_prefixes", [])
    if prefixes:
        return any(href.startswith(p) for p in prefixes)
    return True


# ---------------------------------------------------------------------------
# Link collection
# ---------------------------------------------------------------------------

def _collect_links(
    soup,
    current_url: str,
    root_url: str,
    robots_parser=None,
) -> list[str]:
    """
    Collect all crawlable hrefs from soup.
    Strips <header> and <footer> entirely.
    Strips domain-specific blocked nav elements by ID.
    Passes robots_parser through to _is_allowed_href().
    """
    from bs4 import BeautifulSoup

    soup_copy = BeautifulSoup(str(soup), "html.parser")

    for tag in soup_copy.find_all(["header", "footer"]):
        tag.decompose()

    rules = _get_domain_rules(current_url)
    for blocked_id in rules.get("blocked_nav_ids", []):
        el = soup_copy.find(id=blocked_id)
        if el:
            el.decompose()
            logger.debug(f"  🚫 Removed blocked nav #{blocked_id}")

    seen  = set()
    links = []
    for a in soup_copy.find_all("a", href=True):
        raw = a["href"].split("#")[0].strip()
        if not raw:
            continue
        href = urljoin(current_url, raw).rstrip("/")
        if href in seen:
            continue
        seen.add(href)
        if _is_allowed_href(href, root_url, robots_parser):
            links.append(href)

    return links


# ---------------------------------------------------------------------------
# Screenshot helper
# ---------------------------------------------------------------------------

def _capture_screenshot(driver, name: str):
    '''try:
        timestamp = int(time.time())
        safe_name = "".join(c for c in name if c.isalnum() or c in ("_", "-")).rstrip()
        filepath  = os.path.join(DEBUG_DIR, f"{timestamp}_{safe_name}.png")
        driver.set_window_size(1920, 1080)
        driver.save_screenshot(filepath)
        print(f"📸 Screenshot -> {filepath}")
    except Exception as e:
        logger.warning(f"Screenshot failed: {e}")'''


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def fetch_web_source(
    root_url: str,
    mode: str = "sitemap",
    max_pages: int | None = None,
) -> list[dict]:
    """
    Main entry point for scraping a web source.

    mode:
        "sitemap" (default) — discover the sitemap, fetch all URLs from it.
                              Falls back to BFS crawl if no sitemap found.
                              No page limit — processes every URL in the sitemap.

        "crawl"             — BFS crawl starting from root_url, following links.
                              No page limit by default (max_pages=None = unlimited).
                              Use for sites without sitemaps or when you want
                              to follow dynamic links not listed in the sitemap.

    max_pages:
        Only applies to "crawl" mode. None = no limit (crawl entire site).
        Set to an integer to cap the crawl (useful for testing).
    """
    session = requests.Session()
    session.headers.update(REQUEST_HEADERS)

    print(f"🚀 [{mode.upper()} MODE] Starting: {root_url}")
    content_type, use_selenium = _detect_content_type(session, root_url)

    # PDF root URL — mode doesn't apply, just extract it
    if "application/pdf" in content_type:
        print(f"📋 Target type: PDF")
        return _fetch_pdf_via_selenium(root_url) if use_selenium else _fetch_pdf_via_requests(session, root_url)

    print(f"🌐 Target type: HTML (Selenium: {use_selenium})")

    # Fetch robots.txt once — shared by both modes
    robots        = _fetch_robots_rules(root_url)
    robots_parser = robots["parser"]

    if mode == "sitemap":
        # ── Sitemap mode ──────────────────────────────────────────────────
        sitemap_urls = _discover_sitemap(root_url, robots_parser)
        if sitemap_urls:
            print(f"🗺️  Sitemap mode: {len(sitemap_urls)} URLs to process")
            return _crawl_from_sitemap(session, sitemap_urls, root_url, use_selenium)
        else:
            # No sitemap found — fall back to BFS crawl with no limit
            print(f"⚠️  No sitemap found — falling back to BFS crawl (no page limit)")
            return _crawl_html(
                session, root_url,
                max_pages=None,
                use_selenium=use_selenium,
                robots_rules=robots,
            )

    else:
        # ── BFS crawl mode — no page limit by default ────────────────────
        print(f"🔍 BFS crawl mode — {'unlimited' if max_pages is None else str(max_pages)+' pages max'}")
        return _crawl_html(
            session, root_url,
            max_pages=max_pages,
            use_selenium=use_selenium,
            robots_rules=robots,
        )


def compute_web_fingerprint(pages: list[dict]) -> str:
    if not pages:
        return ""
    parts = sorted(f"{p['url']}:{p['content_hash']}" for p in pages)
    return hashlib.sha256("\n".join(parts).encode()).hexdigest()


# ---------------------------------------------------------------------------
# Content-type detection
# ---------------------------------------------------------------------------

def _detect_content_type(session: requests.Session, url: str) -> tuple[str, bool]:
    try:
        resp = session.head(url, timeout=REQUEST_TIMEOUT, allow_redirects=True)
        if resp.status_code == 200:
            return resp.headers.get("Content-Type", ""), False
        if resp.status_code in (405, 501):
            with session.get(url, stream=True, timeout=REQUEST_TIMEOUT) as r:
                if r.status_code == 200:
                    return r.headers.get("Content-Type", ""), False
    except requests.RequestException:
        pass

    print("🛡️  WAF/Block detected. Initializing headless Selenium browser to bypass...")
    ct = _selenium_get_content_type(url)
    return ct, True


def _selenium_get_content_type(url: str) -> str:
    driver = _make_selenium_driver()
    try:
        print(f"🔍 [Selenium] Probing content type: {url}")
        driver.get(url)
        time.sleep(10)
        _capture_screenshot(driver, "content_type_probe")
        try:
            ct = driver.execute_async_script(
                "const cb = arguments[arguments.length-1];"
                "fetch(arguments[0], {method:'HEAD'})"
                ".then(r => cb(r.headers.get('content-type') || ''))"
                ".catch(() => cb(''))",
                url,
            )
            return ct or ""
        except Exception:
            src = driver.page_source or ""
            if "%PDF" in src or "application/pdf" in src:
                return "application/pdf"
            return "text/html"
    finally:
        driver.quit()

def _get_uc_driver_lock() -> FileLock:
    old_umask = os.umask(0)
    try:
        lock = FileLock(_UC_DRIVER_LOCK_PATH, timeout=90, mode=0o666)
    finally:
        os.umask(old_umask)
    return lock

# ---------------------------------------------------------------------------
# Selenium driver factory
# ---------------------------------------------------------------------------

def _make_selenium_driver(download_dir: str | None = None):
    try:
        import undetected_chromedriver as uc
    except ImportError:
        raise ImportError("undetected-chromedriver is required.")

    options = uc.ChromeOptions()
    options.add_argument("--headless=new")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--window-size=1920,1080")
    options.add_argument(
        "--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
    )

    if download_dir:
        prefs = {
            "download.default_directory":         download_dir,
            "download.prompt_for_download":       False,
            "plugins.always_open_pdf_externally": True,
        }
        options.add_experimental_option("prefs", prefs)
    else:
        # Unique profile per driver instance — prevents SingletonLock conflicts
        # when multiple Chrome sessions start simultaneously (e.g. content-type
        # probe + main crawl). cf_clearance cookie persists within the same
        # driver instance across driver.get() calls without needing a shared dir.
        profile_dir = tempfile.mkdtemp(prefix="chrome_profile_")
        options.add_argument(f"--user-data-dir={profile_dir}")

    chrome_version = getattr(settings, "CHROME_VERSION", None)
    lock = _get_uc_driver_lock()
    with lock:
        driver = uc.Chrome(options=options, version_main=chrome_version)

    driver.set_page_load_timeout(300)
    return driver


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------

def _fetch_pdf_via_requests(session: requests.Session, url: str) -> list[dict]:
    try:
        resp = session.get(url, timeout=REQUEST_TIMEOUT)
        resp.raise_for_status()
    except requests.RequestException as exc:
        logger.error(f"Failed to download PDF {url}: {exc}")
        return []
    return _extract_pdf_pages(io.BytesIO(resp.content), url)


def _fetch_pdf_via_selenium(url: str) -> list[dict]:
    with tempfile.TemporaryDirectory() as tmpdir:
        driver = _make_selenium_driver(download_dir=tmpdir)
        try:
            print(f"📥 [Selenium] Stealth PDF download: {url}")
            driver.get(url)
            pdf_path = None
            for i in range(60):
                files = [f for f in os.listdir(tmpdir) if f.endswith(".pdf")]
                if files:
                    pdf_path = os.path.join(tmpdir, files[0])
                    print(f"⏳ PDF found after {i}s: {pdf_path}")
                    break
                if i == 15:
                    _capture_screenshot(driver, "pdf_download_stalled")
                time.sleep(10)
        finally:
            driver.quit()

        if not pdf_path:
            logger.error(f"PDF did not download within timeout: {url}")
            return []

        with open(pdf_path, "rb") as f:
            return _extract_pdf_pages(io.BytesIO(f.read()), url)


def _extract_pdf_pages(pdf_bytes: io.BytesIO, url: str) -> list[dict]:
    try:
        from pdfminer.high_level import extract_pages
        from pdfminer.layout import LTTextContainer
    except ImportError:
        raise ImportError("pdfminer.six is required. pip install pdfminer.six")

    results: list[dict] = []
    try:
        print(f"⚙️  Parsing PDF: {url}")
        for page_num, page_layout in enumerate(extract_pages(pdf_bytes), start=1):
            parts = [el.get_text() for el in page_layout if isinstance(el, LTTextContainer)]
            text  = "\n".join(parts).strip()
            if not text:
                continue
            print(f"   📖 PDF page {page_num}: {text[:80]}...")
            results.append({
                "url":          url,
                "title":        f"Page {page_num}",
                "text":         text,
                "content_hash": hashlib.sha256(text.encode()).hexdigest(),
                "parent_url":   None,
            })
    except Exception as exc:
        logger.error(f"PDF parse error {url}: {exc}", exc_info=True)

    print(f"✅ PDF done. {len(results)} pages with content.")
    return results


# ---------------------------------------------------------------------------
# DOCX extraction (in-memory, no disk writes)
# ---------------------------------------------------------------------------

def _extract_docx_bytes(docx_bytes: io.BytesIO, url: str) -> list[dict]:
    """
    Extract text from .doc/.docx in-memory using three methods in order:
      1. python-docx  — modern .docx (ZIP + XML format)
      2. mammoth      — old binary .doc (Word 97-2003) + malformed .docx
      3. Raw ZIP      — last resort; strips XML tags from word/document.xml
    Returns [] if all methods fail — never raises, crawl always continues.
    Install: pip install python-docx mammoth
    """
    import zipfile
    import re

    print(f"⚙️  Parsing DOCX in-memory: {url}")

    def _make_result(text: str, method: str) -> list[dict]:
        print(f"   📄 DOCX ({method}): {len(text)} chars from {url}")
        return [{
            "url":          url,
            "title":        url.split("/")[-1][:500],
            "text":         text,
            "content_hash": hashlib.sha256(text.encode()).hexdigest(),
            "parent_url":   None,
        }]

    # ── Method 1: python-docx (modern .docx) ─────────────────────────────
    try:
        from docx import Document
        docx_bytes.seek(0)
        doc   = Document(docx_bytes)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style and para.style.name.startswith("Heading"):
                parts.append(f"\n## {text}\n")
            else:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    c.text.strip() for c in row.cells if c.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        full_text = "\n".join(parts).strip()
        if full_text:
            return _make_result(full_text, "python-docx")
    except ImportError:
        logger.debug("python-docx not installed — pip install python-docx")
    except Exception as e1:
        logger.debug(f"python-docx failed for {url}: {e1}")

    # ── Method 2: mammoth (old .doc binary + some malformed .docx) ───────
    try:
        import mammoth
        docx_bytes.seek(0)
        result    = mammoth.extract_raw_text(docx_bytes)
        full_text = result.value.strip()
        if full_text:
            return _make_result(full_text, "mammoth")
    except ImportError:
        logger.debug("mammoth not installed — pip install mammoth")
    except Exception as e2:
        logger.debug(f"mammoth failed for {url}: {e2}")

    # ── Method 3: raw ZIP/XML extraction (last resort for .docx) ─────────
    try:
        docx_bytes.seek(0)
        with zipfile.ZipFile(docx_bytes, "r") as zf:
            if "word/document.xml" in zf.namelist():
                xml       = zf.read("word/document.xml").decode("utf-8", errors="replace")
                full_text = re.sub(r"\s+", " ", re.sub(r"<[^>]+>", " ", xml)).strip()
                if full_text:
                    return _make_result(full_text, "raw-xml")
    except Exception as e3:
        logger.debug(f"Raw ZIP extraction failed for {url}: {e3}")

    logger.warning(f"⚠️  Could not extract text from {url} — all methods failed, skipping.")
    return []


# ---------------------------------------------------------------------------
# CSV extraction (in-memory, no disk writes)
# ---------------------------------------------------------------------------

def _extract_csv_bytes(csv_bytes: bytes, url: str) -> list[dict]:
    try:
        print(f"⚙️  Parsing CSV in-memory: {url}")
        text_io = io.StringIO(csv_bytes.decode("utf-8", errors="replace"))
        rows    = list(csv_mod.reader(text_io))
        if not rows:
            return []
        parts     = [" | ".join(c.strip() for c in row if c.strip()) for row in rows]
        full_text = "\n".join(p for p in parts if p).strip()
        if not full_text:
            return []
        print(f"   📊 CSV: {len(rows)} rows from {url}")
        return [{
            "url":          url,
            "title":        url.split("/")[-1][:500],
            "text":         full_text,
            "content_hash": hashlib.sha256(full_text.encode()).hexdigest(),
            "parent_url":   None,
        }]
    except Exception as exc:
        logger.error(f"CSV parse error {url}: {exc}", exc_info=True)
        return []


# ---------------------------------------------------------------------------
# HTML text extraction
# ---------------------------------------------------------------------------

def _extract_html_text(soup) -> str:
    from bs4 import BeautifulSoup
    soup_copy = BeautifulSoup(str(soup), "html.parser")
    for tag in soup_copy.find_all(["header", "footer", "script", "style"]):
        tag.decompose()
    content = None
    for selector in CONTENT_SELECTORS:
        content = soup_copy.select_one(selector)
        if content:
            break
    content = content or soup_copy.find("body") or soup_copy
    parts = []
    for tag in content.find_all(["h1", "h2", "h3", "h4", "p", "li", "td", "th"]):
        text = tag.get_text(strip=True)
        if not text:
            continue
        if tag.name in ("h1", "h2", "h3", "h4"):
            parts.append(f"\n## {text}\n")
        else:
            parts.append(text)
    return "\n".join(parts) if parts else soup_copy.get_text(separator="\n", strip=True)


# ---------------------------------------------------------------------------
# Internal sentinel for document redirects
# ---------------------------------------------------------------------------

class _DocumentRedirect(Exception):
    def __init__(self, pages: list[dict], final_url: str):
        self.pages     = pages
        self.final_url = final_url


# ---------------------------------------------------------------------------
# Selenium navigation helpers
# ---------------------------------------------------------------------------

def _fetch_html_via_selenium(
    driver,
    url: str,
    try_click: bool = True,
) -> tuple[str, str, str, object]:
    """
    Navigate to url, wait for CF + page load, simulate reading, then return
    (title, text, final_url, soup).

    Soup is captured AFTER _human_page_interaction() so lazy-loaded links
    that appear after scrolling are included in link discovery.

    Raises _DocumentRedirect if navigation lands on a PDF or DOCX.
    """
    from bs4 import BeautifulSoup

    print(f"🌍 [Selenium] Navigating to: {url}")

    navigated_by_click = False
    if try_click:
        navigated_by_click = _navigate_via_link_click(driver, url)
    if not navigated_by_click:
        driver.get(url)

    _wait_for_cloudflare_if_needed(driver)
    _wait_for_page_ready(driver)

    final_url  = driver.current_url.split("#")[0].rstrip("/")
    final_path = urlparse(final_url).path.lower()
    src_preview = (driver.page_source or "")[:500]

    # ── Detect redirect to PDF ────────────────────────────────────────────
    if final_path.endswith(".pdf") or "%PDF" in src_preview:
        print(f"   📋 Redirect → PDF: {final_url}")
        try:
            _s = requests.Session()
            _s.headers.update(REQUEST_HEADERS)
            r  = _s.get(final_url, timeout=REQUEST_TIMEOUT)
            r.raise_for_status()
            raise _DocumentRedirect(_extract_pdf_pages(io.BytesIO(r.content), final_url), final_url)
        except _DocumentRedirect:
            raise
        except Exception as exc:
            logger.warning(f"PDF redirect fetch failed {final_url}: {exc}")
            return "", "", final_url, None

    # ── Detect redirect to DOCX ───────────────────────────────────────────
    if final_path.endswith((".docx", ".doc")):
        print(f"   📄 Redirect → DOCX: {final_url}")
        try:
            _s = requests.Session()
            _s.headers.update(REQUEST_HEADERS)
            r  = _s.get(final_url, timeout=REQUEST_TIMEOUT)
            r.raise_for_status()
            raise _DocumentRedirect(_extract_docx_bytes(io.BytesIO(r.content), final_url), final_url)
        except _DocumentRedirect:
            raise
        except Exception as exc:
            logger.warning(f"DOCX redirect fetch failed {final_url}: {exc}")
            return "", "", final_url, None

    # ── Simulate reading (scroll) then capture page source ───────────────
    # Page source captured AFTER interaction so lazy-loaded content is included.
    _human_page_interaction(driver)

    url_slug = urlparse(url).path.replace("/", "_")[-30:]
    _capture_screenshot(driver, f"crawl_{url_slug}")

    soup  = BeautifulSoup(driver.page_source, "html.parser")
    title = soup.title.string.strip() if soup.title and soup.title.string else url
    text  = _extract_html_text(soup)

    nav_method = "click" if navigated_by_click else "direct"
    print(f"   📖 [{nav_method}] '{title[:60]}' — {len(text)} chars.")
    return title, text, final_url, soup


def _wait_for_cloudflare_if_needed(driver, max_wait: int = 90) -> bool:
    """
    Wait for CF challenge to clear. Returns True if a challenge was seen.
    Attempts to click the 'I am human' checkbox inside the CF iframe.
    """
    deadline   = time.time() + max_wait
    challenged = False
    while time.time() < deadline:
        title = (driver.title or "").lower()
        src   = driver.page_source or ""
        is_challenge = (
            "just a moment"        in title
            or "checking your browser" in title
            or "attention required"    in title
            or "cf-challenge"          in src
            or "cf_chl_opt"            in src
        )
        if not is_challenge:
            if challenged:
                cool_off = random.uniform(8, 20)
                print(f"✅ CF cleared — cooling off {cool_off:.1f}s")
                time.sleep(cool_off)
            return challenged

        challenged = True
        print("🛡️  Cloudflare challenge detected — waiting…")

        try:
            from selenium.webdriver.common.by import By
            frames = driver.find_elements(By.TAG_NAME, "iframe")
            for frame in frames:
                try:
                    driver.switch_to.frame(frame)
                    cb = driver.find_elements(By.CSS_SELECTOR, 'input[type="checkbox"]')
                    if cb:
                        _human_mouse_move_and_click(driver, cb[0])
                        print("   ☑️  Clicked CF checkbox.")
                    driver.switch_to.default_content()
                    break
                except Exception:
                    driver.switch_to.default_content()
        except Exception:
            pass

        time.sleep(random.uniform(4, 9))

    return challenged


# ---------------------------------------------------------------------------
# Human-behaviour helpers
# ---------------------------------------------------------------------------

def _bezier_curve(p0, p1, p2, p3, steps):
    pts = []
    for i in range(steps + 1):
        t  = i / steps
        mt = 1 - t
        x  = mt**3*p0[0] + 3*mt**2*t*p1[0] + 3*mt*t**2*p2[0] + t**3*p3[0]
        y  = mt**3*p0[1] + 3*mt**2*t*p1[1] + 3*mt*t**2*p2[1] + t**3*p3[1]
        pts.append((x, y))
    return pts


def _human_mouse_move_and_click(driver, element) -> bool:
    try:
        from selenium.webdriver.common.action_chains import ActionChains
        vw    = driver.execute_script("return window.innerWidth")
        vh    = driver.execute_script("return window.innerHeight")
        start = (vw / 2, vh / 2)
        rect  = element.rect
        end   = (
            rect["x"] + rect["width"]  * random.uniform(0.3, 0.7),
            rect["y"] + rect["height"] * random.uniform(0.3, 0.7),
        )
        cp1 = (start[0] + random.uniform(-150, 150), start[1] + random.uniform(-80, 80))
        cp2 = (end[0]   + random.uniform(-100, 100), end[1]   + random.uniform(-60, 60))
        curve   = _bezier_curve(start, cp1, cp2, end, random.randint(18, 32))
        actions = ActionChains(driver)
        actions.move_to_element_with_offset(driver.find_element("tag name", "body"), 0, 0)
        prev = (0.0, 0.0)
        for pt in curve:
            actions.move_by_offset(
                pt[0] - prev[0] + random.gauss(0, 1.5),
                pt[1] - prev[1] + random.gauss(0, 1.0),
            )
            actions.pause(random.uniform(0.008, 0.025))
            prev = pt
        actions.click()
        actions.perform()
        return True
    except Exception as exc:
        logger.debug(f"ActionChains failed ({exc}), JS click fallback.")
        try:
            driver.execute_script("arguments[0].click();", element)
            return True
        except Exception:
            return False


def _navigate_via_link_click(driver, href: str) -> bool:
    try:
        from selenium.webdriver.common.by import By
        candidates = driver.find_elements(By.XPATH, f'//a[@href="{href}"]')
        if not candidates:
            path       = urlparse(href).path
            candidates = driver.find_elements(By.XPATH, f'//a[contains(@href, "{path}")]')
        for el in candidates:
            try:
                if el.is_displayed() and el.is_enabled():
                    driver.execute_script(
                        "arguments[0].scrollIntoView({block:'center',behavior:'smooth'});", el
                    )
                    time.sleep(random.uniform(0.4, 1.2))
                    if _human_mouse_move_and_click(driver, el):
                        return True
            except Exception:
                continue
        return False
    except Exception as exc:
        logger.debug(f"_navigate_via_link_click failed: {exc}")
        return False


def _human_page_interaction(driver):
    try:
        for _ in range(random.randint(*SCROLL_STEPS)):
            driver.execute_script(
                f"window.scrollBy({{top:{random.randint(*SCROLL_AMOUNT)},behavior:'smooth'}});"
            )
            time.sleep(random.uniform(0.6, 2.2))
        if random.random() < 0.35:
            driver.execute_script(
                f"window.scrollBy({{top:-{random.randint(100,350)},behavior:'smooth'}});"
            )
            time.sleep(random.uniform(0.3, 1.0))
    except Exception:
        pass
    time.sleep(
        random.uniform(*DWELL_LONG) if random.random() < DWELL_LONG_PROB
        else random.uniform(*DWELL_SHORT)
    )


def _human_delay_between_pages(pages_fetched: int, crawl_delay: float = 3.0):
    """
    Respect per-domain Crawl-delay from robots.txt (passed in as crawl_delay).
    Every BREAK_EVERY pages take a longer break.
    """
    if pages_fetched > 0 and pages_fetched % BREAK_EVERY == 0:
        pause = random.uniform(*BREAK_SLEEP)
        print(f"☕  Break ({pause:.0f}s) after {pages_fetched} pages…")
        time.sleep(pause)
    else:
        # Use domain's crawl_delay as floor, up to 2x as ceiling
        time.sleep(random.uniform(crawl_delay, crawl_delay * 2))


def _wait_for_page_ready(driver, timeout: int = 30):
    try:
        from selenium.webdriver.support.ui import WebDriverWait
        WebDriverWait(driver, timeout).until(
            lambda d: d.execute_script("return document.readyState") == "complete"
        )
    except Exception:
        time.sleep(5)


def _restart_selenium_session(driver):
    try:
        driver.quit()
    except Exception:
        pass
    print("🔄  Rotating Selenium session…")
    time.sleep(random.uniform(3, 8))
    return _make_selenium_driver()


# ---------------------------------------------------------------------------
# HTML crawler
# ---------------------------------------------------------------------------

def _crawl_html(
    session: requests.Session,
    root_url: str,
    max_pages: int | None = None,   # None = no limit (crawl entire site)
    use_selenium: bool = False,
    delay: float = 1.0,
    robots_rules: dict | None = None,  # pass in if already fetched
) -> list[dict]:
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError("pip install beautifulsoup4 lxml")

    # ── robots.txt rules (reuse if already fetched by caller) ────────────
    if robots_rules is None:
        robots_rules = _fetch_robots_rules(root_url)
    crawl_delay   = robots_rules["crawl_delay"]
    robots_parser = robots_rules["parser"]
    print(f"⏱️  Crawl-delay: {crawl_delay}s | Limit: {'unlimited' if max_pages is None else max_pages}")

    visited: set[str]                    = set()
    queued:  set[str]                    = set()   # URLs in queue but not yet visited
    queue:   list[tuple[str, str|None]]  = [(root_url, None)]
    results: list[dict]                  = []

    driver             = _make_selenium_driver() if use_selenium else None
    pages_fetched      = 0
    session_page_count = 0
    cf_challenge_count = 0   # escalating back-off counter

    # Seed the queued set with the root URL
    queued.add(root_url.rstrip("/"))

    try:
        while queue and (max_pages is None or pages_fetched < max_pages):
            url, parent_url = queue.pop(0)
            url = url.split("#")[0].rstrip("/")

            queued.discard(url)   # remove from queued now that we're processing it

            if not url or url in visited or _should_skip(url):
                continue
            if not _is_allowed_href(url, root_url, robots_parser):
                logger.debug(f"  ⛔ Filtered: {url}")
                continue

            # ── ScrapedURL cache check ────────────────────────────────────
            try:
                from ragbot.models import ScrapedURL
                if ScrapedURL.was_scraped_today(url):
                    logger.info(f"⏭️  Already scraped today: {url}")
                    visited.add(url)
                    continue
            except Exception:
                pass

            visited.add(url)
            print(f"🔍 Queue: {len(queue)} | Fetched: {pages_fetched} | {url}")

            html_source = None
            final_url   = url
            page_hash   = None

            # ── Selenium session rotation ─────────────────────────────────
            if use_selenium and driver and session_page_count >= SESSION_ROTATE_EVERY:
                driver = _restart_selenium_session(driver)
                session_page_count = 0
                cf_challenge_count = 0
                try:
                    driver.get(root_url)
                    _wait_for_cloudflare_if_needed(driver)
                    _wait_for_page_ready(driver)
                    time.sleep(random.uniform(3, 7))
                except Exception:
                    pass

            # ── Fetch ─────────────────────────────────────────────────────
            if use_selenium:
                try:
                    title, text, final_url, soup = _fetch_html_via_selenium(driver, url)

                    # ── CF escalating back-off ────────────────────────────
                    # _wait_for_cloudflare_if_needed returns True if challenged.
                    # We track it via the cool-off print; approximate here by
                    # checking if title/text are empty (CF blocked entirely).
                    if not text and not title:
                        cf_challenge_count += 1
                        print(f"⚠️  CF challenge count this session: {cf_challenge_count}")
                        if cf_challenge_count >= CF_BACKOFF_TRIGGER:
                            backoff = random.uniform(*CF_BACKOFF_SLEEP)
                            print(f"🛑 CF triggered {cf_challenge_count}x — backing off {backoff:.0f}s")
                            time.sleep(backoff)
                            cf_challenge_count = 0
                    else:
                        cf_challenge_count = 0   # reset on successful page

                    if text and text.strip():
                        page_hash = hashlib.sha256(text.encode()).hexdigest()
                        results.append({
                            "url":          final_url,
                            "title":        title[:500],
                            "text":         text,
                            "content_hash": page_hash,
                            "parent_url":   parent_url,
                        })
                    # soup returned from _fetch_html_via_selenium is post-interaction
                    html_source        = soup
                    pages_fetched      += 1
                    session_page_count += 1

                except _DocumentRedirect as doc_redir:
                    results.extend(doc_redir.pages)
                    if doc_redir.pages:
                        page_hash = doc_redir.pages[0]["content_hash"]
                    final_url          = doc_redir.final_url
                    pages_fetched      += 1
                    session_page_count += 1

                except Exception as exc:
                    logger.warning(f"Selenium failed for {url}: {exc}")
                    continue

            else:
                try:
                    print(f"📡 [Requests] Fetching: {url}")
                    resp = session.get(url, timeout=REQUEST_TIMEOUT)

                    if resp.status_code == 403:
                        print(f"🛑 403 — switching to Selenium for {url}")
                        if driver is None:
                            driver = _make_selenium_driver()
                            session_page_count = 0
                        try:
                            title, text, final_url, soup = _fetch_html_via_selenium(
                                driver, url, try_click=False
                            )
                            if text and text.strip():
                                page_hash = hashlib.sha256(text.encode()).hexdigest()
                                results.append({
                                    "url":          final_url,
                                    "title":        title[:500],
                                    "text":         text,
                                    "content_hash": page_hash,
                                    "parent_url":   parent_url,
                                })
                            html_source        = soup
                            pages_fetched      += 1
                            session_page_count += 1
                        except _DocumentRedirect as doc_redir:
                            results.extend(doc_redir.pages)
                            if doc_redir.pages:
                                page_hash = doc_redir.pages[0]["content_hash"]
                            final_url          = doc_redir.final_url
                            pages_fetched      += 1
                            session_page_count += 1
                        except Exception as exc2:
                            logger.warning(f"Selenium fallback failed {url}: {exc2}")
                        if not html_source:
                            continue

                    else:
                        resp.raise_for_status()
                        ct = resp.headers.get("Content-Type", "")

                        if "application/pdf" in ct:
                            pdf_pages = _extract_pdf_pages(io.BytesIO(resp.content), url)
                            results.extend(pdf_pages)
                            if pdf_pages:
                                page_hash = pdf_pages[0]["content_hash"]
                            pages_fetched += 1
                            _human_delay_between_pages(pages_fetched, crawl_delay)
                            try:
                                ScrapedURL.mark_scraped(url, content_hash=page_hash or "")
                            except Exception:
                                pass
                            continue

                        if any(x in ct for x in [
                            "application/vnd.openxmlformats-officedocument.wordprocessingml",
                            "application/msword",
                        ]) or urlparse(url).path.lower().endswith((".docx", ".doc")):
                            docx_pages = _extract_docx_bytes(io.BytesIO(resp.content), url)
                            results.extend(docx_pages)
                            if docx_pages:
                                page_hash = docx_pages[0]["content_hash"]
                            pages_fetched += 1
                            _human_delay_between_pages(pages_fetched, crawl_delay)
                            try:
                                ScrapedURL.mark_scraped(url, content_hash=page_hash or "")
                            except Exception:
                                pass
                            continue

                        if "text/csv" in ct or urlparse(url).path.lower().endswith(".csv"):
                            csv_pages = _extract_csv_bytes(resp.content, url)
                            results.extend(csv_pages)
                            if csv_pages:
                                page_hash = csv_pages[0]["content_hash"]
                            pages_fetched += 1
                            _human_delay_between_pages(pages_fetched, crawl_delay)
                            try:
                                ScrapedURL.mark_scraped(url, content_hash=page_hash or "")
                            except Exception:
                                pass
                            continue

                        if "text/html" not in ct:
                            continue

                        soup  = BeautifulSoup(resp.text, "html.parser")
                        title = soup.title.string.strip() if soup.title and soup.title.string else url
                        text  = _extract_html_text(soup)
                        if text.strip():
                            page_hash = hashlib.sha256(text.encode()).hexdigest()
                            results.append({
                                "url":          url,
                                "title":        title[:500],
                                "text":         text,
                                "content_hash": page_hash,
                                "parent_url":   parent_url,
                            })
                        html_source   = soup
                        pages_fetched += 1

                except requests.RequestException as exc:
                    logger.warning(f"Request failed {url}: {exc}")
                    continue

            # ── Ensure redirect target is also marked visited ────────────
            # Prevents the same PDF/page being fetched again if linked directly
            # by another page under a different URL (e.g. /media/123 → /file.pdf)
            if final_url != url:
                visited.add(final_url)
                queued.discard(final_url)

            # ── Mark as scraped + add final_url to visited ───────────────
            # final_url may differ from url (redirect). Adding it to visited
            # prevents the redirect target being queued and fetched again
            # if another page links to it directly.
            visited.add(final_url)
            queued.discard(final_url)
            if page_hash:
                try:
                    ScrapedURL.mark_scraped(final_url, content_hash=page_hash)
                    if final_url != url:
                        ScrapedURL.mark_scraped(url, content_hash=page_hash)
                except Exception:
                    pass

            # ── Discover child links (post-interaction soup) ──────────────
            if html_source:
                new_links  = _collect_links(html_source, final_url, root_url, robots_parser)
                newly_added = 0
                for href in new_links:
                    if href not in visited and href not in queued:
                        queue.append((href, final_url))
                        queued.add(href)
                        newly_added += 1
                print(f"  🔗 {newly_added} new links queued from {final_url} (queue size: {len(queue)})")

            # ── Inter-page delay (respects Crawl-delay) ───────────────────
            _human_delay_between_pages(pages_fetched, crawl_delay)

    finally:
        if driver:
            print("🔌 Closing Selenium session.")
            driver.quit()

    print(f"✅ Crawl complete. {len(results)} pages from {root_url} ({pages_fetched} URLs fetched)")
    logger.info(f"Crawled {len(results)} pages from {root_url}")
    return results